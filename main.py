import streamlit as st
import os
from dotenv import load_dotenv
from langchain_openai import ChatOpenAI
from langchain.schema import HumanMessage, AIMessage
from langchain.memory import ConversationBufferMemory
from langchain.chains import ConversationChain
from langchain.prompts import PromptTemplate
from langchain.callbacks.base import BaseCallbackHandler
from typing import Any, List
from phoenix.otel import register

# RAG-related imports
import faiss
import numpy as np
import PyPDF2
import docx
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain.embeddings.base import Embeddings
from langchain_community.vectorstores import FAISS
from langchain.chains import RetrievalQA
from langchain.schema import Document
import requests
import json
import pickle
import tempfile

# tracer_provider = register(
#   project_name="Nvidia Chatbot",
#   endpoint="http://localhost:6006/v1/traces",
#   auto_instrument=True
# )

# Load environment variables
load_dotenv()

class NVIDIAEmbeddings(Embeddings):
    """Custom embedding class for NVIDIA NeMo Retriever embedding model"""
    
    def __init__(self, api_key: str = None, model: str = None):
        self.api_key = api_key or os.getenv("EMBEDDING_NVIDIA_API_KEY")
        self.model = model or os.getenv("EMBEDDING_MODEL_NAME", "nvidia/llama-3.2-nemoretriever-300m-embed-v1")
        self.base_url = "https://integrate.api.nvidia.com/v1"
    
    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        """Embed a list of documents"""
        embeddings = []
        for text in texts:
            embedding = self._get_embedding(text, input_type="passage")
            if embedding:
                embeddings.append(embedding)
        return embeddings
    
    def embed_query(self, text: str) -> List[float]:
        """Embed a single query"""
        return self._get_embedding(text, input_type="query")
    
    def _get_embedding(self, text: str, input_type: str = "passage") -> List[float]:
        """Get embedding for a single text"""
        if not self.api_key:
            st.error("EMBEDDING_NVIDIA_API_KEY not found in environment variables")
            return []
        
        if not self.model:
            st.error("EMBEDDING_MODEL_NAME not found in environment variables")
            return []
        
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
        
        # Truncate text if too long (NVIDIA API has limits)
        if len(text) > 4000:
            text = text[:4000]
        
        payload = {
            "model": self.model,
            "input": [text],  # NVIDIA API expects an array
            "input_type": input_type,  # Required for asymmetric models: "passage" for documents, "query" for queries
            "encoding_format": "float"
        }
        
        try:
            response = requests.post(
                f"{self.base_url}/embeddings",
                headers=headers,
                json=payload,
                timeout=30
            )
            
            if response.status_code != 200:
                st.error(f"API Error {response.status_code}: {response.text}")
                return []
            
            result = response.json()
            return result["data"][0]["embedding"]
            
        except requests.exceptions.RequestException as e:
            st.error(f"Request error: {str(e)}")
            return []
        except Exception as e:
            st.error(f"Error getting embedding: {str(e)}")
            return []

class DocumentProcessor:
    """Handle processing of different document types"""
    
    @staticmethod
    def extract_text_from_pdf(file) -> str:
        """Extract text from PDF file"""
        try:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
        except Exception as e:
            st.error(f"Error processing PDF: {str(e)}")
            return ""
    
    @staticmethod
    def extract_text_from_docx(file) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            st.error(f"Error processing DOCX: {str(e)}")
            return ""
    
    @staticmethod
    def extract_text_from_txt(file) -> str:
        """Extract text from TXT file"""
        try:
            return file.read().decode('utf-8')
        except Exception as e:
            st.error(f"Error processing TXT: {str(e)}")
            return ""
    
    @staticmethod
    def process_uploaded_file(uploaded_file) -> str:
        """Process uploaded file based on its type"""
        if uploaded_file is None:
            return ""
        
        file_extension = uploaded_file.name.split('.')[-1].lower()
        
        if file_extension == 'pdf':
            return DocumentProcessor.extract_text_from_pdf(uploaded_file)
        elif file_extension == 'docx':
            return DocumentProcessor.extract_text_from_docx(uploaded_file)
        elif file_extension == 'txt':
            return DocumentProcessor.extract_text_from_txt(uploaded_file)
        else:
            st.error(f"Unsupported file type: {file_extension}")
            return ""

class VectorStoreManager:
    """Manage FAISS vector store operations"""
    
    def __init__(self, embeddings: NVIDIAEmbeddings):
        self.embeddings = embeddings
        self.vector_store = None
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
        )
    
    def create_vector_store(self, documents: List[Document]) -> FAISS:
        """Create a new FAISS vector store from documents"""
        if not documents:
            st.warning("No documents provided to create vector store")
            return None
        
        try:
            # Split documents into chunks
            texts = self.text_splitter.split_documents(documents)
            
            # Create vector store
            self.vector_store = FAISS.from_documents(texts, self.embeddings)
            return self.vector_store
            
        except Exception as e:
            st.error(f"Error creating vector store: {str(e)}")
            return None
    
    def add_documents(self, documents: List[Document]):
        """Add new documents to existing vector store"""
        if self.vector_store is None:
            return self.create_vector_store(documents)
        
        try:
            # Split documents into chunks
            texts = self.text_splitter.split_documents(documents)
            
            # Add to existing vector store
            self.vector_store.add_documents(texts)
            
        except Exception as e:
            st.error(f"Error adding documents: {str(e)}")
    
    def search_documents(self, query: str, k: int = 4) -> List[Document]:
        """Search for relevant documents"""
        if self.vector_store is None:
            return []
        
        try:
            docs = self.vector_store.similarity_search(query, k=k)
            return docs
        except Exception as e:
            st.error(f"Error searching documents: {str(e)}")
            return []
    
    def save_vector_store(self, path: str):
        """Save vector store to disk"""
        if self.vector_store is not None:
            try:
                self.vector_store.save_local(path)
            except Exception as e:
                st.error(f"Error saving vector store: {str(e)}")
    
    def load_vector_store(self, path: str):
        """Load vector store from disk"""
        try:
            self.vector_store = FAISS.load_local(path, self.embeddings, allow_dangerous_deserialization=True)
        except Exception as e:
            st.error(f"Error loading vector store: {str(e)}")
            self.vector_store = None


class StreamlitCallbackHandler(BaseCallbackHandler):
    """Custom callback handler for streaming responses to Streamlit"""
    
    def __init__(self, container):
        self.container = container
        self.text = ""
    
    def on_llm_new_token(self, token: str, **kwargs: Any) -> None:
        """Handle new token from LLM"""
        self.text += token
        self.container.markdown(self.text + "▌")
    
    def on_llm_end(self, response, **kwargs: Any) -> None:
        """Handle end of LLM response"""
        self.container.markdown(self.text)

class NVIDIAChatbot:
    def __init__(self):
        self.api_key = os.getenv("NVIDIA_API_KEY")
        self.model_name = os.getenv("MODEL_NAME")
        self.app_title = os.getenv("APP_TITLE", "LangChain Chatbot")
        
        # Initialize LangChain ChatOpenAI with NVIDIA endpoint
        self.llm = ChatOpenAI(
            model=self.model_name,
            openai_api_key=self.api_key,
            openai_api_base="https://integrate.api.nvidia.com/v1",
            temperature=0.7,
            max_tokens=1024,
            streaming=True
        )
        
        # Initialize conversation memory
        self.memory = ConversationBufferMemory(
            return_messages=True,
            memory_key="history"
        )
        
        # Create conversation prompt template
        self.prompt_template = PromptTemplate(
            input_variables=["history", "input"],
            template="""You are a helpful AI assistant powered by NVIDIA's technology. 
            Be conversational, helpful, and engaging in your responses.

            Chat History:
            {history}

            Human: {input}
            AI Assistant:"""
        )
        
        # RAG-specific prompt template
        self.rag_prompt_template = PromptTemplate(
            input_variables=["context", "question"],
            template="""You are a helpful AI assistant that answers questions based on the provided context. 
            Use the context below to answer the question. If the answer cannot be found in the context, 
            say "I cannot find that information in the provided documents."

            Context:
            {context}

            Question: {question}
            
            Answer:"""
        )
        
        # Initialize conversation chain
        self.conversation_chain = ConversationChain(
            llm=self.llm,
            memory=self.memory,
            prompt=self.prompt_template,
            verbose=False
        )
        
        # Initialize RAG components
        self.embeddings = NVIDIAEmbeddings()
        self.vector_store_manager = VectorStoreManager(self.embeddings)
        self.rag_chain = None
        
    def create_rag_chain(self):
        """Create RAG chain when vector store is available"""
        if self.vector_store_manager.vector_store is not None:
            self.rag_chain = RetrievalQA.from_chain_type(
                llm=self.llm,
                chain_type="stuff",
                retriever=self.vector_store_manager.vector_store.as_retriever(search_kwargs={"k": 4}),
                chain_type_kwargs={"prompt": self.rag_prompt_template}
            )
        
    def add_documents_to_knowledge_base(self, documents: List[Document]):
        """Add documents to the knowledge base"""
        self.vector_store_manager.add_documents(documents)
        self.create_rag_chain()
        
    def get_rag_response(self, query: str, container):
        """Get response using RAG"""
        if self.rag_chain is None:
            container.markdown("❌ No documents loaded for RAG. Please upload documents first.")
            return None
            
        try:
            # Create callback handler for streaming
            callback_handler = StreamlitCallbackHandler(container)
            
            # Get response from RAG chain
            response = self.rag_chain.run(
                query=query,
                callbacks=[callback_handler]
            )
            
            return response
            
        except Exception as e:
            st.error(f"Error in RAG response: {str(e)}")
            return None
    
    def validate_config(self):
        """Validate that required environment variables are set"""
        if not self.api_key:
            st.error("NVIDIA_API_KEY not found in environment variables. Please set it in your .env file.")
            st.stop()
        
        if not self.model_name:
            st.error("MODEL_NAME not found in environment variables. Please set it in your .env file.")
            st.stop()
    
    def update_model(self, new_model_name: str):
        """Update the model and reinitialize the LLM"""
        if new_model_name != self.model_name:
            self.model_name = new_model_name
            
            # Reinitialize LLM with new model
            self.llm = ChatOpenAI(
                model=self.model_name,
                openai_api_key=self.api_key,
                openai_api_base="https://integrate.api.nvidia.com/v1",
                temperature=0.7,
                max_tokens=1024,
                streaming=True
            )
            
            # Reinitialize conversation chain with new LLM
            self.conversation_chain = ConversationChain(
                llm=self.llm,
                memory=self.memory,
                prompt=self.prompt_template,
                verbose=False
            )
    
    def get_streaming_response(self, user_input: str, container):
        """Get streaming response using LangChain"""
        try:
            # Create callback handler for streaming
            callback_handler = StreamlitCallbackHandler(container)
            
            # Get response from conversation chain with streaming
            response = self.conversation_chain.predict(
                input=user_input,
                callbacks=[callback_handler]
            )
            
            return response
            
        except Exception as e:
            st.error(f"Error calling NVIDIA API: {str(e)}")
            return None
    
    def get_conversation_history(self):
        """Get formatted conversation history"""
        messages = []
        if hasattr(self.memory, 'chat_memory') and self.memory.chat_memory.messages:
            for message in self.memory.chat_memory.messages:
                if isinstance(message, HumanMessage):
                    messages.append({"role": "user", "content": message.content})
                elif isinstance(message, AIMessage):
                    messages.append({"role": "assistant", "content": message.content})
        return messages
    
    def clear_memory(self):
        """Clear conversation memory"""
        self.memory.clear()

def initialize_session_state():
    """Initialize Streamlit session state"""
    if "messages" not in st.session_state:
        st.session_state.messages = []
    
    if "chatbot" not in st.session_state:
        st.session_state.chatbot = NVIDIAChatbot()
    
    if "chat_mode" not in st.session_state:
        st.session_state.chat_mode = "regular"  # "regular" or "rag"
        
    if "uploaded_documents" not in st.session_state:
        st.session_state.uploaded_documents = []

def main():
    st.set_page_config(
        page_title="NVIDIA RAG Chatbot",
        page_icon="🤖",
        layout="wide"
    )
    
    # Initialize session state
    initialize_session_state()
    
    # Validate configuration
    st.session_state.chatbot.validate_config()
    
    # App header
    st.title("🤖 NVIDIA RAG Chatbot")
    st.markdown("Powered by NVIDIA NIM API, LangChain, and FAISS Vector Store")
    
    # Chat mode selector
    col1, col2 = st.columns([3, 1])
    with col1:
        chat_mode = st.radio(
            "Chat Mode:",
            ["Regular Chat", "Document Q&A (RAG)"],
            index=0 if st.session_state.chat_mode == "regular" else 1,
            horizontal=True
        )
        st.session_state.chat_mode = "regular" if chat_mode == "Regular Chat" else "rag"
    
    # Sidebar for configuration and document management
    with st.sidebar:
        st.header("⚙️ Configuration")
        
        # Model selection
        current_model = st.text_input(
            "Model Name",
            value=st.session_state.chatbot.model_name,
            help="Enter the NVIDIA model name (e.g., meta/llama3-70b-instruct)"
        )
        
        # Update model if changed
        if current_model != st.session_state.chatbot.model_name:
            st.session_state.chatbot.update_model(current_model)
            st.success(f"Model updated to: {current_model}")
        
        st.markdown("---")
        
        # Document Management Section
        st.header("📄 Document Management")
        
        # File uploader
        uploaded_files = st.file_uploader(
            "Upload Documents",
            type=['pdf', 'docx', 'txt'],
            accept_multiple_files=True,
            help="Upload PDF, DOCX, or TXT files for document Q&A"
        )
        
        if uploaded_files:
            if st.button("📥 Process Documents"):
                with st.spinner("Processing documents..."):
                    documents = []
                    for uploaded_file in uploaded_files:
                        # Extract text from file
                        text = DocumentProcessor.process_uploaded_file(uploaded_file)
                        
                        if text.strip():
                            # Create document object
                            doc = Document(
                                page_content=text,
                                metadata={"source": uploaded_file.name}
                            )
                            documents.append(doc)
                            st.session_state.uploaded_documents.append(uploaded_file.name)
                    
                    if documents:
                        # Add documents to knowledge base
                        st.session_state.chatbot.add_documents_to_knowledge_base(documents)
                        st.success(f"✅ Processed {len(documents)} documents!")
                    else:
                        st.error("❌ No text could be extracted from the uploaded files.")
        
        # Show uploaded documents
        if st.session_state.uploaded_documents:
            st.subheader("📚 Loaded Documents")
            for doc_name in st.session_state.uploaded_documents:
                st.text(f"📄 {doc_name}")
            
            if st.button("🗑️ Clear All Documents"):
                st.session_state.uploaded_documents = []
                st.session_state.chatbot.vector_store_manager.vector_store = None
                st.session_state.chatbot.rag_chain = None
                st.success("✅ All documents cleared!")
                st.rerun()
        
        st.markdown("---")
        
        # Chat controls
        st.header("💬 Chat Controls")
        
        # Clear chat button
        if st.button("🔄 Clear Chat History"):
            st.session_state.chatbot.clear_memory()
            st.session_state.messages = []
            st.rerun()
        
        # Show memory info
        memory_messages = st.session_state.chatbot.get_conversation_history()
        st.info(f"Messages in memory: {len(memory_messages)}")
        
        st.markdown("---")
        
        # Mode-specific info
        if st.session_state.chat_mode == "rag":
            st.header("🔍 RAG Mode")
            if st.session_state.chatbot.vector_store_manager.vector_store is not None:
                st.success("✅ Knowledge base ready")
                st.markdown("""
                **RAG Features:**
                - Document-based Q&A
                - Semantic search
                - Context-aware responses
                - Source attribution
                """)
            else:
                st.warning("⚠️ Upload documents to enable RAG")
        else:
            st.header("💬 Regular Chat")
            st.markdown("""
            **Chat Features:**
            - General conversation
            - Memory-based context
            - Streaming responses
            - NVIDIA NIM powered
            """)
        
        # Display current model
        st.info(f"🤖 Current Model: {st.session_state.chatbot.model_name}")
    
    # Display chat messages
    if st.session_state.chat_mode == "rag":
        memory_messages = []  # Don't show regular chat history in RAG mode
    else:
        memory_messages = st.session_state.chatbot.get_conversation_history()
    
    for message in memory_messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])
    
    # Chat input
    prompt_placeholder = "Ask questions about your documents..." if st.session_state.chat_mode == "rag" else "What would you like to know?"
    
    if prompt := st.chat_input(prompt_placeholder):
        # Display user message
        with st.chat_message("user"):
            st.markdown(prompt)
        
        # Generate assistant response
        with st.chat_message("assistant"):
            message_placeholder = st.empty()
            
            # Choose response method based on mode
            if st.session_state.chat_mode == "rag":
                if st.session_state.chatbot.vector_store_manager.vector_store is None:
                    message_placeholder.markdown("❌ Please upload documents first to use RAG mode.")
                else:
                    with st.spinner("Searching documents..."):
                        response = st.session_state.chatbot.get_rag_response(prompt, message_placeholder)
            else:
                with st.spinner("Thinking..."):
                    response = st.session_state.chatbot.get_streaming_response(prompt, message_placeholder)
            
            if not response:
                st.error("Failed to get response. Please check your configuration.")


if __name__ == "__main__":
    main()
